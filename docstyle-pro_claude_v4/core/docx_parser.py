"""
docx_parser.py — 원본 .docx 를 분석해 구조화 JSON 으로 변환

지원 블록 타입:
  chapter_header / h1 / h2 / h3 / body
  quote_box / tip_box / insight_box / caution_box / warning_box
  qa_block / bullet / image_ref / table / hr / empty
"""

from __future__ import annotations
import os
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


# ── 박스 패턴 감지 키워드 ──────────────────────────────────

BOX_PATTERNS: dict[str, list[str]] = {
    "tip_box":     ["Tip ", "TIP ", "팁 "],
    "caution_box": ["주의 ", "주의:", "⚠ ", "Warning"],
    "insight_box": ["인사이트", "핵심 포인트", "핵심:", "결론:"],
    "warning_box": ["⚠  의학적", "⚠  법적", "고지:", "경고:"],
    "quote_box":   ["AI 답변", "\"", """],
}

PHASE_PATTERN = re.compile(r"^\[?\\?\[?Phase\s*\d+", re.IGNORECASE)
NUMBER_PATTERN = re.compile(r"^(\d+)\.\s+(.+)")


def _detect_box_type(text: str) -> str | None:
    for box_type, keywords in BOX_PATTERNS.items():
        if any(text.strip().startswith(kw) for kw in keywords):
            return box_type
    return None


def _get_heading_level(para) -> str | None:
    style_name = para.style.name if para.style else ""
    if "Heading 1" in style_name:
        return "h1"
    if "Heading 2" in style_name:
        return "h2"
    if "Heading 3" in style_name:
        return "h3"
    # 휴리스틱: 볼드 + 폰트 크기
    for r in para.runs:
        size = r.font.size
        if r.bold and size:
            pt = size.pt
            if pt >= 16:
                return "h1"
            if pt >= 14:
                return "h2"
            if pt >= 13:
                return "h3"
    return None


def _has_image(para) -> bool:
    return para._element.find(qn("w:drawing")) is not None


def _extract_table(table) -> dict:
    if not table.rows:
        return {}
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    rows = [
        [cell.text.strip() for cell in row.cells]
        for row in table.rows[1:]
    ]
    return {"type": "table", "headers": headers, "rows": rows}


# ── 메인 파서 ──────────────────────────────────────────────

def parse_docx(docx_path: str) -> dict:
    """
    Returns:
    {
        "meta": { "title": str, "author": str, "source_file": str },
        "images": [ { "id": str, "filename": str, ... } ],
        "content": [ { "type": str, ... } ]
    }
    """
    doc = Document(docx_path)
    content: list[dict] = []
    image_index = 0

    # ── 단락 순회 ──────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()

        # 이미지 참조
        if _has_image(para):
            image_index += 1
            content.append({
                "type":     "image_ref",
                "image_id": f"image{image_index}",
                "caption":  text,
            })
            continue

        if not text:
            continue

        # 챕터 헤더 ([Phase N])
        if PHASE_PATTERN.match(text):
            content.append({
                "type":  "chapter_header",
                "phase": text,
                "title": "",
                "subtitle": "",
            })
            continue

        # 제목 레벨
        heading = _get_heading_level(para)
        if heading:
            match = NUMBER_PATTERN.match(text)
            if heading == "h1" and match:
                content.append({
                    "type":   "h1",
                    "number": match.group(1),
                    "text":   match.group(2),
                })
            else:
                content.append({"type": heading, "text": text})
            continue

        # 박스 타입
        box_type = _detect_box_type(text)
        if box_type:
            content.append({"type": box_type, "text": text})
            continue

        # 불릿 목록
        if para.style and "List" in para.style.name:
            content.append({"type": "bullet", "text": text})
            continue

        # 일반 본문
        content.append({"type": "body", "text": text})

    # ── 표 순회 ───────────────────────────────────────────
    for table in doc.tables:
        table_block = _extract_table(table)
        if table_block:
            content.append(table_block)

    # ── 메타 추출 ─────────────────────────────────────────
    core = doc.core_properties
    meta = {
        "title":       core.title or Path(docx_path).stem,
        "author":      core.author or "",
        "source_file": os.path.basename(docx_path),
    }

    return {"meta": meta, "content": content}


def parse_and_save(docx_path: str, output_json: str) -> dict:
    """파싱 후 JSON 파일로 저장"""
    result = parse_docx(docx_path)
    Path(output_json).parent.mkdir(parents=True, exist_ok=True)
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    return result


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python docx_parser.py <input.docx> <output.json>")
        sys.exit(1)
    parse_and_save(sys.argv[1], sys.argv[2])
    print(f"[OK] → {sys.argv[2]}")
